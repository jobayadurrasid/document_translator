from docx import Document
from deep_translator import GoogleTranslator
import time

# Load the original document
input_file = "target.docx"
output_file = "expected.docx"

doc = Document(input_file)
translated_doc = Document()

translator = GoogleTranslator(source='auto', target='en')

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        translated_doc.add_paragraph("")
        continue

    retries = 10
    success = False
    while retries > 0 and not success:
        try:
            print(f"Translating paragraph {i+1}/{len(doc.paragraphs)}...")
            translated_text = translator.translate(text)
            translated_doc.add_paragraph(translated_text)
            success = True
        except Exception as e:
            retries -= 1
            print(f"[Retry {10 - retries}/10] Error translating paragraph {i+1}: {e}")
            time.sleep(2)  # wait before retry

    if not success:
        print(f"❌ Skipped paragraph {i+1} after 10 failed attempts.")

# Save the translated document
translated_doc.save(output_file)
print(f"\n✅ Translation complete. Saved to: {output_file}")
